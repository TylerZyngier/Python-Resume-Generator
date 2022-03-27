# BEFORE RUNNING APP MAKE SURE TO RUN COMMAND: pip install -r requirements.txt

from docx import Document
from docx.shared import Inches

document = Document()

# name, phone number, and email details
name = input('Your name: ')
phone_number = input('Your phone number: ')
email = input('Your email: ')

document.add_paragraph(
    name + '\n' + phone_number + '\n' + email)

# about me
document.add_heading('About Me')
document.add_paragraph(
    input('\nTell me about yourself: ')
)

# work experience
document.add_heading("Work Experience")


def add_experience():
    p = document.add_paragraph()

    company = input('\nEnter a company you have worked for: ')
    from_date = input('From date: ')
    to_date = input('To date: ')

    p.add_run(company + ' ').bold = True
    p.add_run(from_date + '-' + to_date + '\n').italic = True

    experience_details = input(
        'Describe your experience at ' + company + ': ')
    p.add_run(experience_details)


add_experience()

while True:
    has_more_experience = input(
        'Do you have more work experience? (Yes or No) ')

    if has_more_experience.lower() == 'yes':
        add_experience()
    elif has_more_experience.lower() == 'no':
        break
    else:
        print('Invalid answer, please input Yes or No')
        continue

# skills
document.add_heading('Skills')


def add_skill():
    skill = input('Enter a skill you have: ')
    p = document.add_paragraph(skill)
    p.style = 'List Bullet'


print('\n')
add_skill()

while True:
    has_more_skills = input(
        'Do you have more skills? (Yes or No) ')

    if has_more_skills.lower() == 'yes':
        add_skill()
    elif has_more_skills.lower() == 'no':
        break
    else:
        print('Invalid answer, please input Yes or No')
        continue

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "Resume generated with Python Resume Generator Project by Tyler Zyngier"

# save the document
document.save('resume.docx')
